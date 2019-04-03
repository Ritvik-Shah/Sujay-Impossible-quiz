from PIL import Image
import PIL
from docx import Document
import os

img = Image.open('Sujay.jpg')
img_dude = img.resize((100, 100), PIL.Image.ANTIALIAS)
img_dude.save('Meme_Sujay.jpg')
document = Document()

p1 = document.add_paragraph('When you pass the Sujay test, you be like')
paragraph = document.add_picture('Meme_Sujay.jpg')
document.save('Sujay.docx')

Sujay = input("Are you Sujay? ")
if Sujay == 'no' or Sujay == 'No':                                                                                
    img = Image.open('Sujay.jpg')
    img.show()
    os.remove('Sujay.docx')
else:
    print("You will now have to take test to see if you are really Sujay")
    q1 = input("What is your true name? ")
    if q1 == 'Big Mac' or q1 == 'big mac' or q1 == 'Sanjay' or q1 == 'sanjay' or q1 == 'God' or q1 == 'god':
        print("Well done. You have answered question 1 correctly")
        q2 = input("What is your favorite catch phrase? ")
        if q2 == 'You are stupid' or q2 == 'You are wrong' or q2 == 'Nitish is Dumb' or q2 == 'Aquaman and Wonderwoman were amazing' or q2 == 'Sexy time' or q2 == 'I am black':
            print("Wow, you are the real " + q1 + ". I am sorry that I ever doubted you. Forgive me " +q1+ " because " +q2)
            print("Here is your prize!")
            yeet = input("Are you satisfied with the quiz? ")
            if yeet == 'no' or yeet == 'No':
                img = Image.open('Sujay.jpg')
                img.show()
                yeet2 = input("Are you satisfied now? ")
                if yeet2 == 'yes' or yeet2 == 'Yes':
                    img2 = Image.open('Sujay full.jpg')
                    img2.show()
                    os.system('start Sujay.docx')
                    remove = input("Would you like to remove the " + q1+ " document? ")
                    if remove == 'yes' or remove == 'Yes':
                        os.remove('Sujay.docx')
                    else:
                        os._exit(0)
                else:
                    print("How have you failed this quiz")
                    print("You are the dumbest person I know")
                    os.remove('Sujay.docx')
            else:
                print("Liar!!!!, You are not the real " + q1)
                os.remove('Sujay.docx')
        else:
            print("Liar!!!!, You are not the real " + q1)
            os.remove('Sujay.docx')
    else:
        print("Liar!!!!, You are not the real " + q1)
        os.remove('Sujay.docx')
